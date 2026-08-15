"""Fabrication du .docx avec python-docx. Word refusait les fichiers produits par
la bibliothèque JavaScript ; celle-ci écrit un format plus conservateur."""
import json, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLEU = RGBColor(0x1F, 0x38, 0x64)
GRIS = RGBColor(0x66, 0x66, 0x66)

items = json.load(open(sys.argv[1], encoding="utf8"))
sortie = sys.argv[2]
titre  = sys.argv[3] if len(sys.argv) > 3 else "Document"

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
for s in doc.sections:
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2)
    p = s.header.paragraphs[0]; p.text = titre
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.size = Pt(8); p.runs[0].font.color.rgb = RGBColor(0x99,0x99,0x99)
    q = s.footer.paragraphs[0]; q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = q.add_run(); r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99,0x99,0x99)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    q._p.append(fld)

def par(txte="", taille=12, gras=False, ital=False, coul=None, avant=0, apres=6,
        just=True, retrait=None, saut=False, style=None):
    p = doc.add_paragraph(style=style)
    if saut: p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(avant)
    p.paragraph_format.space_after = Pt(apres)
    if just: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if retrait: p.paragraph_format.left_indent = Pt(retrait)
    r = p.add_run(txte)
    r.font.size = Pt(taille); r.bold = gras; r.italic = ital
    if coul is not None: r.font.color.rgb = coul
    return p

def tableau(head, rows):
    t = doc.add_table(rows=1, cols=len(head))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(head):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(str(h)); r.bold = True; r.font.size = Pt(9)
        sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), "1F3864")
        c._tc.get_or_add_tcPr().append(sh)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row[:len(head)]):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(9)
    doc.add_paragraph()

COUL = {"rouge": "8E1B1B", "orange": "9C5A05", "vert": "1C5E36", "gris": "3C4557"}
PRIO = {"bloquant": "8E1B1B", "critique": "9C5A05", "important": "1F3864", "information": "767676"}

def fond(p, hexa):
    sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), hexa)
    p._p.get_or_add_pPr().append(sh)

def bordure_gauche(p, hexa, taille="18"):
    pb = OxmlElement("w:pBdr"); b = OxmlElement("w:left")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), taille)
    b.set(qn("w:color"), hexa); b.set(qn("w:space"), "6")
    pb.append(b); p._p.get_or_add_pPr().append(pb)

def cadre(p, hexa):
    pb = OxmlElement("w:pBdr")
    for cote in ("top", "left", "bottom", "right"):
        b = OxmlElement("w:" + cote)
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), hexa); b.set(qn("w:space"), "5")
        pb.append(b)
    p._p.get_or_add_pPr().append(pb)

for i in items:
    k = i.get("k")
    if k == "bandeau":
        c = COUL.get(i.get("couleur"), "3C4557")
        p = par(i["t"], 22, gras=True, coul=RGBColor(0xFF, 0xFF, 0xFF),
                avant=10, apres=2, just=False)
        fond(p, c)
        q = par(i["sous"], 10, coul=RGBColor(0xFF, 0xFF, 0xFF), apres=10, just=False)
        fond(q, c)
        continue
    if k == "etape":
        t = i["t"] + (f"   ({i['compte']})" if i.get("compte") else "")
        p = par(t, 12, gras=True, coul=BLEU, avant=12, apres=5, just=False)
        pb = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "10")
        b.set(qn("w:color"), "1F3864"); b.set(qn("w:space"), "3")
        pb.append(b); p._p.get_or_add_pPr().append(pb)
        continue
    if k == "acte":
        c = PRIO.get(i.get("priorite"), "1F3864")
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"{i['n']}. "); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLEU
        r = p.add_run(i["t"]); r.bold = True; r.font.size = Pt(11)
        r = p.add_run(f"   [{i['priorite']}]"); r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(c)
        bordure_gauche(p, c)
        tete = f"{i['etat']} — " if i.get("etat") else ""
        q = par(f"{tete}{i['pourquoi']}  ·  {i['id']}", 9, coul=GRIS, apres=6)
        bordure_gauche(q, c)
        continue
    if k == "saut":
        doc.add_paragraph().paragraph_format.page_break_before = True
        continue
    if k == "interdit":
        c = {"certain": "8E1B1B", "reserve": "9C5A05", "examen": "767676"}.get(i.get("ton"), "8E1B1B")
        p = par(i["t"], 11.5, gras=True, coul=RGBColor.from_string(c),
                avant=8, apres=1, just=False)
        cadre(p, c)
        q = par(f"{i['pourquoi']}  ·  {i['id']}", 9, coul=GRIS, apres=8)
        cadre(q, c)
        continue
    if k == "acquis":
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Pt(14); p.paragraph_format.first_line_indent = Pt(-14)
        r = p.add_run("✓  "); r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1C, 0x5E, 0x36)
        r = p.add_run(i["t"]); r.font.size = Pt(10)
        r = p.add_run(" — " + i["base"]); r.font.size = Pt(10); r.font.color.rgb = GRIS
        continue
    if k == "sur":   par(i["t"], 10, coul=GRIS, apres=3)
    elif k == "t1":  par(i["t"], 20, gras=True, coul=BLEU, apres=4, just=False)
    elif k == "trait":
        p = par("", apres=8); pb = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6"); b.set(qn("w:color"), "BFBFBF")
        b.set(qn("w:space"), "4"); pb.append(b); p._p.get_or_add_pPr().append(pb)
    elif k == "h1":  par(i["t"], 16, gras=True, coul=BLEU, avant=0, apres=8, just=False, saut=True)
    elif k == "h2":  par(i["t"], 13, gras=True, coul=BLEU, avant=12, apres=6, just=False)
    elif k == "h3":  par(i["t"], 12, gras=True, coul=BLEU, avant=10, apres=5, just=False)
    elif k == "p":   par(i["t"])
    elif k == "note":par(i["t"], 10, coul=GRIS)
    elif k == "puce":par("—  " + i["t"], retrait=14)
    elif k == "enc":
        par(i["titre"], 12, gras=True, coul=BLEU, avant=10, apres=3, just=False)
        par(i["t"], 11, ital=True, apres=10)
    elif k == "table": tableau(i["head"], i["rows"])
    elif k == "piece":
        par(f"Pièce n° {i['num']} — {i['titre']}", 13, gras=True, coul=BLEU, avant=14, apres=3, just=False)
        par(f"Nature : {i['nature']} · Émetteur : {i['emetteur']} · Date : {i['date']}", 9, coul=GRIS, apres=2)
        par(f"Ce qu'elle prouve : {i['prouve']}" + (f" · Fondement : {i['texte']}" if i.get("texte") else ""), 9, coul=GRIS, apres=6)
    elif k == "doc":
        for l in i["lignes"]:
            if isinstance(l, str): par(l, 11, retrait=14)
            elif l.get("k") == "c": par(l["t"], 11, gras=True, just=False)
            elif l.get("k") == "r": par(l["t"], 11, just=False)
            elif l.get("k") == "table": tableau(l.get("head") or [""]*len(l["rows"][0]), l["rows"])
            else: par(l.get("t", ""), 11, retrait=14)
    elif k == "sign": par(i["t"], 10, ital=True, coul=GRIS, just=False, apres=10)
    elif k == "schema":
        par(i["titre"], 12, gras=True, coul=BLEU, avant=12, apres=2, just=False)
        par(i["article"], 10, ital=True, coul=GRIS, apres=4, just=False)
        head = [c["t"] for c in i["cols"]]
        rows = []
        for rang in i["bandes"]:
            ligne = []
            for b in rang: ligne += [b["t"]] + [""] * (b["span"] - 1)
            rows.append(ligne)
        tableau(head, rows)

doc.save(sortie)
print("Word :", sortie)
