#!/usr/bin/env python3
"""Générateur de documents Word, à partir d'une liste d'éléments en JSON.

RÈGLE DU DÉPÔT : tous les .docx se produisent avec python-docx, jamais avec la
bibliothèque JavaScript « docx ». Les fichiers produits par celle-ci sont
refusés par Microsoft Word — « Impossible d'ouvrir le fichier Office Open XML.
Des problèmes ont été décelés dans son contenu. » python-docx écrit un format
plus conservateur, que Word accepte.

Éléments reconnus, dans l'ordre où ils sont donnés :
  sur    surtitre, gris                    {"k":"sur","t":"…"}
  t1     titre du document                 {"k":"t1","t":"…"}
  trait  filet horizontal                  {"k":"trait"}
  h1 h2 h3   titres de niveau              {"k":"h1","t":"…"}
  p      paragraphe                        {"k":"p","t":"…"}
  note   paragraphe gris, plus petit       {"k":"note","t":"…"}
  puce   liste à puces                     {"k":"puce","t":"…"}
  enc    encadré titré                     {"k":"enc","titre":"…","t":"…"}
  table  tableau                           {"k":"table","head":[…],"rows":[[…]]}
  saut   saut de page                      {"k":"saut"}

Usage : python3 word_py.py elements.json sortie.docx
"""
import json
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

BLEU = RGBColor(0x1F, 0x38, 0x64)
GRIS = RGBColor(0x5F, 0x68, 0x74)
ENCRE = RGBColor(0x16, 0x18, 0x1D)


def _bord(par):
    """Un filet sous le paragraphe — le trait de séparation."""
    p = par._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    bas = OxmlElement("w:bottom")
    bas.set(qn("w:val"), "single")
    bas.set(qn("w:sz"), "6")
    bas.set(qn("w:space"), "4")
    bas.set(qn("w:color"), "DCDFE4")
    bd.append(bas)
    p.append(bd)


def _cadre(par, fond="F4F6FB"):
    """Un fond et un cadre : l'encadré."""
    p = par._p.get_or_add_pPr()
    ombre = OxmlElement("w:shd")
    ombre.set(qn("w:val"), "clear")
    ombre.set(qn("w:fill"), fond)
    p.append(ombre)
    bd = OxmlElement("w:pBdr")
    for cote in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + cote)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "6")
        e.set(qn("w:color"), "1F3864")
        bd.append(e)
    p.append(bd)


def _par(doc, texte, taille=11, gras=False, couleur=None, avant=0, apres=6, style=None):
    par = doc.add_paragraph(style=style)
    par.paragraph_format.space_before = Pt(avant)
    par.paragraph_format.space_after = Pt(apres)
    r = par.add_run(texte or "")
    r.font.size = Pt(taille)
    r.bold = gras
    if couleur is not None:
        r.font.color.rgb = couleur
    return par


def construire(elements, chemin):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for e in elements:
        k = e.get("k")
        if k == "sur":
            _par(doc, e.get("t"), taille=9, couleur=GRIS)
        elif k == "t1":
            _par(doc, e.get("t"), taille=20, gras=True, couleur=ENCRE, avant=6, apres=10)
        elif k == "trait":
            _bord(_par(doc, "", taille=1, apres=8))
        elif k == "h1":
            _par(doc, e.get("t"), taille=15, gras=True, couleur=BLEU, avant=14, apres=6)
        elif k == "h2":
            _par(doc, e.get("t"), taille=13, gras=True, couleur=BLEU, avant=12, apres=5)
        elif k == "h3":
            _par(doc, e.get("t"), taille=11.5, gras=True, couleur=BLEU, avant=10, apres=4)
        elif k == "p":
            _par(doc, e.get("t"))
        elif k == "note":
            _par(doc, e.get("t"), taille=9.5, couleur=GRIS)
        elif k == "puce":
            _par(doc, e.get("t"), style="List Bullet", apres=3)
        elif k == "enc":
            _cadre(_par(doc, e.get("titre"), gras=True, couleur=BLEU, avant=10, apres=2))
            _cadre(_par(doc, e.get("t"), apres=10))
        elif k == "saut":
            doc.add_page_break()
        elif k == "table":
            head = e.get("head") or []
            rows = e.get("rows") or []
            t = doc.add_table(rows=1, cols=len(head))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(head):
                cel = t.rows[0].cells[i]
                cel.text = ""
                r = cel.paragraphs[0].add_run(str(h))
                r.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                sh = OxmlElement("w:shd")
                sh.set(qn("w:val"), "clear")
                sh.set(qn("w:fill"), "1F3864")
                cel._tc.get_or_add_tcPr().append(sh)
            for ligne in rows:
                cells = t.add_row().cells
                for i, v in enumerate(ligne[: len(head)]):
                    cells[i].text = ""
                    r = cells[i].paragraphs[0].add_run("" if v is None else str(v))
                    r.font.size = Pt(9.5)
            _par(doc, "", taille=6, apres=6)
        else:
            raise ValueError("élément inconnu : %r" % (k,))

    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2)
        s.top_margin = s.bottom_margin = Cm(2.0)

    doc.save(chemin)
    return chemin


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(__doc__)
        sys.exit(2)
    with open(sys.argv[1], encoding="utf-8") as f:
        elements = json.load(f)
    construire(elements, sys.argv[2])
    print("%s — %d éléments" % (sys.argv[2], len(elements)))
